"""SharePoint connector -- the fifth and last of Milestone 9's document/chat
connectors (PROJECT_PLAN.md section 14), following `teams.py`: both use
Microsoft Graph API with an already-issued bearer access token.

Implements `app.ingestion.connectors.base.Connector` structurally (a
`Protocol` -- see that module's docstring for why no explicit inheritance is
used).

Expected `ResolvedConnectorConfig.config` shape for this source:
    {"site_ids": ["<graph-site-id-1>", "<graph-site-id-2>", ...]}
Graph *site IDs* (not site URLs/names) -- the same "IDs the IT Admin already
knows" choice every other multi-container connector built this milestone
makes (`JiraConnector`'s projects, `AzureDevOpsConnector`'s projects,
`ConfluenceConnector`'s spaces).

Uses Graph's drive **delta** endpoint (`/sites/{id}/drive/root/delta`) to
list a site's default document library recursively in one flat, paginated
walk -- unlike a plain `/children` listing, which only returns one folder
level at a time and would need this connector to walk the folder tree
itself. Known limitation, flagged rather than silently built around: delta
sync is *meant* to be driven by a persisted `@odata.deltaLink` token across
syncs (so a later call only sees what changed) -- this connector does not
yet persist that token anywhere (the `Connector` protocol's `cursor` only
lives for one `fetch_batch` sequence, not across separate sync runs), so
every sync -- full or incremental -- re-walks the entire delta from
scratch, exactly like `GitHubConnector._list_tree_page`'s equivalent,
already-documented tradeoff for full-tree syncs. `since` is instead applied
as a client-side filter on `lastModifiedDateTime`, the same fallback
`TeamsConnector` already uses for the same underlying reason (no
`$filter`-style query support on this endpoint).

Plain-text files, plus Word/PDF/Excel documents (`_EXTRACTORS`), have their
content fetched and ingested; folders and every other file type (images,
PowerPoint, ...) are listed by the delta walk but skipped, not erroring --
the same "skipped, not an error" treatment `GitHubConnector._fetch_file_
content` already gives binary/undecodable files. `.pdf`/`.docx`/`.xlsx`
extraction uses `pypdf`/`python-docx`/`openpyxl` respectively -- each is a
best-effort text join (PDF: every page's `extract_text()`; DOCX: every
paragraph; XLSX: every non-empty cell, read-only mode so a large workbook
isn't fully loaded into memory) with no layout/formatting preserved, since
this connector's output is ingestible search content, not a faithful
document rendering. A file that fails to parse (corrupt, password-
protected, actually a different format than its extension claims, ...) is
skipped exactly like an undecodable plain-text file already is -- one bad
file must not fail the whole sync.

Known limitation, flagged rather than silently built around: delta sync is
*meant* to be driven by a persisted `@odata.deltaLink` token across syncs
(so a later call only sees what changed) -- this connector does not yet
persist that token anywhere (the `Connector` protocol's `cursor` only lives
for one `fetch_batch` sequence, not across separate sync runs), so every
sync -- full or incremental -- re-walks the entire delta from scratch,
exactly like `GitHubConnector._list_tree_page`'s equivalent, already-
documented tradeoff for full-tree syncs. `since` is instead applied as a
client-side filter on `lastModifiedDateTime`, the same fallback
`TeamsConnector` uses for its own comparable gap.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

import docx
import httpx
from openpyxl import load_workbook
from pypdf import PdfReader

from app.ingestion.schemas import FetchResult, RawDocument, ResolvedConnectorConfig
from app.shared.config.logging import get_logger

logger = get_logger(__name__)

_GRAPH_API_BASE_URL = "https://graph.microsoft.com/v1.0/"
_PLAIN_TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}


def _extract_pdf_text(raw_bytes: bytes) -> str:
    """Join every page's extracted text -- see module docstring's
    "best-effort text join, no layout preserved" note.
    """
    reader = PdfReader(BytesIO(raw_bytes))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(raw_bytes: bytes) -> str:
    """Join every paragraph's text, in document order."""
    document = docx.Document(BytesIO(raw_bytes))
    return "\n\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def _extract_xlsx_text(raw_bytes: bytes) -> str:
    """Join every non-empty cell's string value, row by row, sheet by
    sheet -- `read_only=True` so a large workbook streams rather than
    loading fully into memory (see module docstring).
    """
    workbook = load_workbook(BytesIO(raw_bytes), read_only=True, data_only=True)
    try:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value) for value in row if value is not None]
                if cells:
                    lines.append("\t".join(cells))
        return "\n".join(lines)
    finally:
        workbook.close()


# Extension -> extractor, dispatched by `_fetch_text_content`. Plain-text
# extensions are handled separately (a plain decode, not one of these
# binary-format parsers) -- see that method.
_EXTRACTORS = {
    ".pdf": _extract_pdf_text,
    ".docx": _extract_docx_text,
    ".xlsx": _extract_xlsx_text,
}


@dataclass
class _SharePointClient:
    """What `SharePointConnector.authenticate` returns as
    `AuthenticatedClient`. Bundles the authenticated HTTP client with the
    fixed site-ID list -- `fetch_batch` only receives this object back, not
    the original config, same reasoning as `_TeamsClient`.
    """

    http: httpx.AsyncClient
    site_ids: list[str]


class SharePointConnector:
    """Fetches plain-text and Office/PDF files from a fixed set of
    SharePoint sites' default document libraries -- see module docstring
    for the exact supported-format list.
    """

    source_name = "sharepoint"
    # Same "no single published flat ceiling" situation `TeamsConnector`
    # documents for Graph API in general. 1.0 req/s is a conservative
    # steady-state budget for the delta-listing + per-file content-download
    # calls this connector makes.
    requests_per_second = 1.0

    async def authenticate(self, config: ResolvedConnectorConfig) -> _SharePointClient:
        """Build an authenticated Graph client from `config`.

        `config.credential_ref` is treated as a literal, already-issued
        bearer access token -- same flagged placeholder as
        `TeamsConnector.authenticate`. Calls `GET /me` once so an invalid/
        expired token fails loudly here rather than on the first real
        fetch.
        """
        http = httpx.AsyncClient(
            base_url=_GRAPH_API_BASE_URL,
            headers={"Authorization": f"Bearer {config.credential_ref}"},
            timeout=30.0,
        )
        try:
            response = await http.get("me")
            response.raise_for_status()
        except Exception:
            await http.aclose()
            logger.warning(
                "sharepoint_authenticate_failed",
                connector_config_id=str(config.connector_config_id),
            )
            raise

        site_ids = list(config.config.get("site_ids", []))
        logger.info(
            "sharepoint_authenticate_succeeded",
            connector_config_id=str(config.connector_config_id),
            site_count=len(site_ids),
        )
        return _SharePointClient(http=http, site_ids=site_ids)

    async def fetch_batch(
        self,
        client: _SharePointClient,
        *,
        since: datetime | None,
        cursor: str | None,
    ) -> FetchResult:
        """Fetch one page of raw SharePoint drive items from
        `client.site_ids`.

        `cursor` is this connector's own opaque JSON envelope
        `{"site_index": int, "next_link": str | None}` -- `next_link` is
        Graph's own full `@odata.nextLink` URL, the same "resume mid-list
        via Graph's own opaque absolute URL" shape `TeamsConnector`'s cursor
        already uses.

        Only file-type items with a supported text extension and a fetched
        content body end up in the returned page -- folders, unsupported
        file types, and content-download failures are filtered out here,
        inside `fetch_batch`, not in `normalize` (which is sync-only and
        cannot itself perform the content-download network call -- the same
        "do the I/O in `fetch_batch`, attach the result onto the raw item"
        shape `GitHubConnector._list_file_items_page` already uses for its
        own per-file content fetch).
        """
        site_index, next_link = self._decode_cursor(cursor)

        if site_index >= len(client.site_ids):
            return FetchResult(items=[], next_cursor=None, has_more=False)

        site_id = client.site_ids[site_index]

        if next_link is not None:
            response = await client.http.get(next_link)
        else:
            response = await client.http.get(f"sites/{site_id}/drive/root/delta")
        response.raise_for_status()
        payload = response.json()

        raw_entries: list[dict[str, Any]] = payload.get("value", [])
        items: list[dict[str, Any]] = []
        for entry in raw_entries:
            if "file" not in entry:
                continue  # a folder, or a deleted-item tombstone -- not ingestible content
            if since is not None and not self._is_recent_enough(entry, since):
                continue
            content = await self._fetch_text_content(client.http, entry)
            if content is None:
                continue  # unsupported extension, or download/decode failure -- skipped
            entry["_site_id"] = site_id
            entry["_content"] = content
            items.append(entry)

        graph_next_link = payload.get("@odata.nextLink")

        if graph_next_link:
            next_state = {"site_index": site_index, "next_link": graph_next_link}
            has_more = True
        else:
            # No `@odata.nextLink` means Graph returned `@odata.deltaLink`
            # instead -- this site's delta walk is complete for this sync.
            next_site_index = site_index + 1
            next_state = {"site_index": next_site_index, "next_link": None}
            has_more = next_site_index < len(client.site_ids)

        return FetchResult(
            items=items,
            next_cursor=json.dumps(next_state) if has_more else None,
            has_more=has_more,
        )

    async def _fetch_text_content(
        self, http: httpx.AsyncClient, entry: dict[str, Any]
    ) -> str | None:
        """Download and extract one file's text content, or `None` if it's
        not a supported extension or fails to decode/parse -- see module
        docstring's supported-format list and "skipped, not an error" note.
        """
        name = entry.get("name", "")
        extension = next(
            (ext for ext in (*_PLAIN_TEXT_EXTENSIONS, *_EXTRACTORS) if name.endswith(ext)),
            None,
        )
        if extension is None:
            return None
        download_url = entry.get("@microsoft.graph.downloadUrl")
        if not download_url:
            return None
        response = await http.get(download_url)
        try:
            response.raise_for_status()
        except Exception:
            return None

        if extension in _PLAIN_TEXT_EXTENSIONS:
            try:
                return response.text
            except Exception:
                # Any decode failure is treated the same as "not ingestible
                # text", not a reason to fail the whole batch -- same
                # "skipped, not an error" treatment `GitHubConnector.
                # _fetch_file_content` gives its own undecodable files.
                return None

        try:
            return _EXTRACTORS[extension](response.content)
        except Exception:
            # A corrupt/password-protected/mislabeled file must not fail
            # the whole batch -- same treatment as an undecodable text file.
            logger.warning(
                "sharepoint_content_extraction_failed", name=name, extension=extension
            )
            return None

    @staticmethod
    def _is_recent_enough(entry: dict[str, Any], since: datetime) -> bool:
        """Client-side `since` filter -- see module docstring's "Known
        limitations" note. Entries with an unparseable/missing
        `lastModifiedDateTime` are kept, not silently dropped, the same
        "absence of a timestamp isn't evidence of staleness" reasoning
        `TeamsConnector._is_recent_enough` already uses.
        """
        raw_timestamp = entry.get("lastModifiedDateTime")
        if not raw_timestamp:
            return True
        try:
            modified_at = datetime.fromisoformat(raw_timestamp.replace("Z", "+00:00"))
        except ValueError:
            return True
        if modified_at.tzinfo is None:
            modified_at = modified_at.replace(tzinfo=timezone.utc)
        return modified_at >= since

    def normalize(self, raw_item: Any) -> RawDocument:
        """Convert one raw SharePoint drive item dict (with `"_site_id"`/
        `"_content"` injected by `fetch_batch`) into a `RawDocument`.
        """
        site_id = raw_item["_site_id"]
        item_id = raw_item["id"]
        name = raw_item.get("name", item_id)
        content = raw_item["_content"]

        metadata: dict[str, str] = {"site_id": site_id}
        if raw_item.get("lastModifiedDateTime"):
            metadata["updated"] = raw_item["lastModifiedDateTime"]
        parent_path = (raw_item.get("parentReference") or {}).get("path")
        if parent_path:
            metadata["folder_path"] = parent_path

        return RawDocument(
            source=self.source_name,
            external_id=f"{site_id}:{item_id}",
            content=content,
            title=name,
            source_url=raw_item.get("webUrl"),
            metadata=metadata,
        )

    async def close(self, client: _SharePointClient) -> None:
        """Close the underlying `httpx.AsyncClient` opened by `authenticate`."""
        await client.http.aclose()

    @staticmethod
    def _decode_cursor(cursor: str | None) -> tuple[int, str | None]:
        """Parse this connector's opaque cursor envelope back into
        `(site_index, next_link)`, defaulting to the first site with no
        Graph-native `next_link` when `cursor` is None (a full sync's first
        page, or an incremental sync's first page for this run).
        """
        if cursor is None:
            return 0, None
        state = json.loads(cursor)
        return int(state["site_index"]), state.get("next_link")
